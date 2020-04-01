# -*- coding: cp1252 -*-
#TO EXTRACT NAMES OF CHARACTERS(WHICH ARE IN BOLD) FROM .docx file
from docx import *
import time
import numpy as np
from scipy.stats import zscore
import itertools
import collections
plot = 1
start = time.time()
document = Document("D:\\ML\\Character_Matrix\\Ramayana.docx")  #COMMENT: Opens Word Document
bolds=[]
for para in document.paragraphs:
    for run in para.runs:
        if run.bold :
            word=run.text
            if word!="":
                if word[0].isupper():
                    
                    bolds.append(word)
#print(*bolds,sep="\n")
    
sno=1
count=0  
character=dict()
unique=set(bolds)
#print(unique)   
for i in unique:
    for j in bolds:
        if i==j:
            count+=1
    if i!="":
        character[i]=count #COMMENT counting the frequency of each word?
    count=0;
    
ranked_list=sorted(character, key=lambda x: character[x])[::-1]
for i in ranked_list:
    #print(sno,")",i,character[i])
    sno+=1



#print(ranked_list[0:14])

#TO NAVIGATE WORDS AND SENTENCES IN .TXT FILE WHICH HAS THE SAME CONTENTS AS THE .DOCX FILE
f=open("D:\\ML\\Character_Matrix\\Ramayana.txt","r")
text=f.read()
sentences=text.split(".")
relation=[]#contains a list of tuples of characters appearing together
rel=[]#contains two characters that appear together
c=0
for sentence in sentences:
    words=sentence.split()
    for i in range(0,len(words)):
        #Creates list of two characters appearing togetheR        
        if words[i] in bolds:
            rel.append(words[i])
        else:
            continue
        for j in range(i+1,len(words)):
            if words[j] in bolds:
                if(len(rel)<2 and words[j]!=rel[0]):#Createst the final pair
                    rel.append(words[j])
                    relation.append(tuple(rel))		#COMMENT: what is this doing??
                    rel=[words[i]]
        rel=[]#Reinitializing for next pair
                
                
key_set=[]
relationship=dict()
for i in range(0, len(relation)):
    item=relation[i]
    count=1
    if(set(item) not in key_set):
        for j in range(i+1, len(relation)):
            if(set(item)==set(relation[j])):
                count+=1
        key_set.append(set(item))
        relationship[item]=count #COMMENT: number of times the relationship has occured
        count=0
'''       
#RANKING IN ORDER OF IMPORTANCE OF RELATIONSHIP
ranked_list=sorted(relationship, key=lambda x: relationship[x])[::-1]
#print(ranked_list)
n = [i for i in unique]
for i in ranked_list:
    if(n_ in i):
        print(i,relationship[i])

'''
f.close()







document = Document("D:\\ML\\Character_Matrix\\Ramayana.docx")
bolds=[]
for para in document.paragraphs:
    for run in para.runs:
        if run.bold :
            word=run.text
            if word!="":
                if word[0].isupper():
                    
                    bolds.append(word) #COMMENT: Assuming that every word that starts with a capital letter is a proper noun?

count=0  
character=dict()
unique=set(bolds)
#print(unique)   
for i in unique:
    for j in bolds:
        if i==j:
            count+=1
    if i!="":
        character[i]=count
    count=0;
    
bolds=sorted(character, key=lambda x: character[x])[::-1]
#print(bolds)
male_pronouns=["he","his","He","His","him","Him","sons","twins","twin"]
female_pronouns=["she","hers","her","She","Hers","Her","mother","Mother","wife"]

f=open("D:\\ML\\Character_Matrix\\Ramayana.txt","r")
text=f.read()
sentences=text.split(".")
male=dict()
female=dict()
mcount=0
fcount=0
k=0
for i in bolds:
    #print(i)
    for j in sentences:
        words=j.split()
        for word in words:
            if word==i:   #COMMENT: why compare word and i? number of pronouns in proximity of bold word?
                for gender in words:
                    if gender in male_pronouns:
                        #print("MALE")
                        mcount+=1
                    elif gender in female_pronouns:
                        #print("FEMALE")
                        fcount+=1
                            
        
    """if i=="Kusha":
        print(mcount)
        print(fcount)
    """
    male[i]=mcount
    female[i]=fcount
    mcount=0
    fcount=0

final=dict()
males=[]
females=[]
ambiguous=[]
for i in bolds:
    #print(i)
    if male[i]>female[i]:
        final[i]="M"
        males.append(i)
    elif female[i]>male[i]:
        final[i]="F"
        females.append(i)
    else:
        ambiguous.append(i)
#COMMENT: Classifying the proper nouns. 
monkeys = list()
monkeys = ['Vali','Hanuman','Sugriva','Kabandha']









graph = dict()
graph_transitive = dict()
for char in unique:
    graph[char]=dict()
    graph_transitive[char]=dict()
    for char_ in unique:
        graph[char][char_] = 0
        graph_transitive[char][char_] = 0
        
        
for c1,c2 in relation:
    try:
        graph[c1][c2]+=1
        graph[c2][c1]+=1
        graph_transitive[c1][c2] = 1
        graph_transitive[c2][c1] = 1 #COMMENT: What's this for?
    except Exception as e:
        print(e)
        pass

graph_transitive_clousure = graph_transitive



for k in (graph):
    for i in (graph):
        for j in (graph):
            try:
                graph_transitive_clousure[i][j] = graph_transitive_clousure[i][j] or (graph_transitive_clousure[i][k] and graph_transitive_clousure[k][j]) 
            except:
                print(k,i,j)




new_graph = dict()
for x,y in relation:
    if x not in new_graph:
        new_graph[x]=list()
    if not y in new_graph[x]:
        new_graph[x].append(y)
for char in unique:
    if char not in new_graph:    
        new_graph[char] = list()



for i in unique:
    graph_transitive_clousure[i][i]=0



if (plot):
    length = 1804 # fix dis nigga
    plotx = []
    for x in graph:
        for y in graph:
            plotx.append(graph[x][y])
    plotx.sort(reverse = True)
    ploty = list(range(1,len(plotx)+1))
    import math
    from matplotlib import pyplot as plt
    from scipy.stats import pearsonr
    plotxLog = list(map(lambda x: math.log(x),plotx[:-length]))
    plotyLog = list(map(lambda x: math.log(x),ploty[:-length]))
    plt.plot(ploty[:-length],plotx[:-length])
    plt.xlabel('Relationship Rank')
    plt.ylabel('Strength of relation')
    plt.show()
    plt.plot(plotyLog,plotxLog)
    plt.xlabel('Log(Relationship Rank)')
    plt.ylabel('Log(Strength of relation)')
    plt.show()
    corr, _ = pearsonr(plotxLog, plotyLog)
    print("Alpah",abs(corr)) 
    



recurssive_counter = 0
no_of_chars = 5
def dfs (graph,char,char_init,visited = list()):
    global recurssive_counter
    global no_of_chars
    recurssive_counter+=1
    if graph_transitive_clousure[char][char_init] and recurssive_counter > no_of_chars:
        recurssive_counter=0
        return visited
    if (char in visited):
        recurssive_counter=0
        return visited
    else:
        visited.append(char)
        for ch in graph[char] :
            dfs(graph,ch,char_init,visited)
    recurssive_counter=0
    return visited


def std(c1,c2):
    visited = dfs (new_graph,c1,c2)    
    print(len(visited))

    summary = list()
    for i in range(len(visited) - 1) :
        summary.append((visited[i],visited[i+1]))
    
    final_sollution = list()
    for x,y in summary:
        final_sollution.append(graph[x][y])
    
    final_sollution = np.asarray(final_sollution,dtype = np.int8)
    print('std:',np.std(final_sollution))
    

def avg():
    result = []
    for i in unique:
        for j in unique:
            if not i == j:
               result.append(graph[i][j])
    return 2*sum(result)/len(result)


def new_std (char_list):
    if len(char_list) == 2:
            if graph[char_list[0]][char_list[1]] > 20:
                return 10.0
            else:
                return 30.0
    no_of_chars = len(char_list)
    result = list()
    char = list()
    normalize = list()
    if no_of_chars <= 1 :
        return [0]
    else:
        for i in char_list:
            for j in char_list:
                if i == j :
                    continue
                char.append((i,j))
        #print(char)
        for x,y in char:
            result.append(graph[x][y])
        result.sort()
        result = list(set(result))
        result = np.asarray(result,dtype = np.int8)
        std = np.std(result)
        z_socre = zscore(result)
        mean = np.mean(result)
        for i in result:
            normalize.append((i - mean)/z_socre)
        #normalize = list(map(lambda x : (x-mean)/z_score,result))
        #print(normalize,round(np.std(normalize)),round(np.mean(normalize)),sep = '\n')
        if np.std(normalize) == 0.0:
            return 100.0
        return abs(np.std(normalize))
    
#val = new_std(['Jatayu','Sita','Hanuman'])

#thresh = 1/3500
#val=val*(thresh-val)/100
#val = 1/((1/thresh)*val)

#data = 1/((thresh*(val-10)**2)+1)
#print('STF',abs(val))




#print(time.time() - start)   


def find_max(graph,char):
    maxy = 0
    result = None
    for c1 in graph:
        if graph[char][c1] > maxy:
            maxy = graph[char][c1]
            result = c1
    return result
            


def predict_graph(chars,probs,thresh = 0.15):
    orig_std = new_std(chars)
    if orig_std < 25.0: # its decent
        print("Perfect")
        return chars
    chars_temp = list(map(lambda x:x , chars))
    good_bad = list(map(lambda x: x > thresh ,probs))
    char_prob = list(zip(chars,good_bad))
    good_predictions = list()
    #print(char_prob,chars_temp)
    for c,p in char_prob:
        if not p:
            chars_temp.remove(c)
            temp_std = new_std(chars_temp)
            #print(temp_std,chars_temp)
            if temp_std < orig_std:
                good_predictions.append((chars_temp,temp_std)) 
        else:
            continue
    temp_list = list()
    #print(len(chars) - len(good_predictions),len(chars))
    
    if len(good_predictions) == 0 and len(chars_temp) > 2:
        print('Decent')
        return chars_temp
    
    if len(good_predictions) == 0:
        for x,y in char_prob:
            if not y:
                char_prob.remove((x,y))
                #print(char_prob)
                good_predictions_temp = list(map(lambda x:x, char_prob))
                new_char = list(map(lambda x: x[0], good_predictions_temp))
                good_predictions = list()
                good_predictions.append((new_char,20))
                good_predictions.sort(key = lambda x: x [1])
                print('hf',good_predictions)                
    else:
        print()
    if len(good_predictions) < len(chars):
        for new_chars,_ in good_predictions:
            miny = orig_std
            #print(new_chars)
            for char in list(unique):
                print(new_chars)
                temp_list = list(map(lambda x:x , new_chars))
                gay = list(map(lambda x:x , new_chars))
                gay.append(char)
                temp_std = new_std(gay)
                if temp_std <= orig_std:
                    #miny = temp_std
                    temp_list.append(char)
                    temp_list=list(set(temp_list))
                    print(temp_list)
        return temp_list
    else:
        print("Big Fak cannot fix dis shise")
        return chars
    

    
def trans (char_list):
    combinations = list(set(list(itertools.combinations(char_list,2))))
    for x,y in combinations:
        if not graph_transitive_clousure[x][y]: 
            return False
    return True
    

def new_graph (char_list, prob_list):
    #char_list = list(set(char_list))
    if type(char_list) == 'NoneType':
        print('Type Error')
        return []
    if len(char_list) <= 1:
        return char_list
    if len(char_list) == 2:
        '''
        if new_std(char_list) < 20.0: std
            return char_list
        '''
        if graph_transitive_clousure[char_list[0]][char_list[1]]:
            return char_list 
        else:
            if prob_list[0] > prob_list[1]:
                return [char_list[0]]
            else:
                return [char_list[1]]
    prob_list = list(zip(char_list,list(map(lambda x : x > 0.15 , prob_list))))
    char_list = list(set(char_list))
    prob_list.sort(key = lambda x : x[1])
    
    removed_list = list(map(lambda x : x[0],(filter(lambda x: x[1], prob_list))))
    best_prediction_graph = dict()
    print('After removing :',removed_list)
    for char in removed_list:
        best_prediction_graph[char]=(find_max(graph,char))
    best_replacements = list(set(best_prediction_graph.values()))
    print("Beest replacement",best_replacements)
    possible_stds = list()
    for char in best_replacements:
        possible_list = list(map(lambda x:x , removed_list))
        possible_list.append(char)
        if not trans(possible_list):
            continue
        possible_std = new_std(possible_list)
        possible_stds.append((possible_std,possible_list))
    possible_stds.sort(key = lambda x: x[0])
    if len(possible_stds) == 0:
        return list(set(char_list))
    else:
        return list(set(possible_stds[0][1]))
    #print(best_replacements)
    
#print('Final ',new_graph(['Vali','Shabari','Rama'],[0.1,0.7,0.5]))
       
def in_vector (char):
    result = 0
    for c in unique:
        result += graph[c][char]
    return result

def out_vector (char):
    result = 0
    for c in unique:
        result += graph[char][c]
    return result

def total_vectors():
    result = 0
    for char in unique:
        result = result + out_vector(char) + in_vector(char)
    return result/4.0

'''
def find_pairs():
    for char in unique:
        if (abs(in_vector(char) - out_vector(char)) / 1.0 ) :
            
'''         
    
def find_list_gender (graph,char,gender): # the best gender for give n charecter
    if gender == 'female':
        return sorted(females,key = lambda x : graph[x][char])[::-1]
    elif gender == 'male':
        return sorted(males,key = lambda x : graph[x][char])[::-1]
    elif gender == 'monkeys':
        return sorted(monkeys,key = lambda x : graph[x][char])[::-1]
    else:
        return sorted(ambiguous,key = lambda x : graph[x][char])[::-1]
    
    
def get_gender (char):  # finds gender of charecter
    if char in females:
        return 'female'
    elif char in monkeys:
        return 'monkeys'
    elif char in males:
        return 'male'
    else:
        return 'idk'
 
    
    
    
def val_map(value, istart, istop, ostart, ostop):
  return ostart + (ostop - ostart) * ((value - istart) / (istop - istart))



    
def find_std (char_list):  ## return true / false based on std
    if len(char_list) <= 2: return True
    else:
        pairs = list()
        for i in char_list:
            for j in char_list:
                if not i == j:
                    pairs.append((i,j))
        #vals = list(map(lambda x:graph[x[0]][x[1]],pairs))
        vals = list()
        #print(pairs)
        for x,y in pairs:
            vals.append(graph[x][y])
        vals = np.asarray(vals,dtype = np.int8)
        #print("max(vals): ",max(vals),"min(vals):", min(vals))
        try:
            new_val = list(map(lambda x: val_map(x,min(vals),max(vals),0.0,10.0),vals))
        except Exception as e:
            print('Python Error',e)
            return True
        #print(np.std(new_val))        
        if np.std(new_val) < 4.2:
            return True
        else:
            return False
        
    
def remove_duplicates (char_list_gender,count=0):  # removes duplicates and replaces it with best fit
    for char,gender in char_list_gender:
        if get_gender(char) == gender: continue
        else: char_list_gender[char_list_gender.index((char,gender))] = (char,get_gender(char))
    if len(char_list_gender) == len(set(char_list_gender)):
        return char_list_gender
    #print(char_list_gender)
    duplicate = [item for item, count in collections.Counter(char_list_gender).items() if count > 1]
    remaining = [item for item, count in collections.Counter(char_list_gender).items() if count == 1]
    for t in duplicate:
        remaining.append(t)
    #print(remaining)
    #print(duplicate)
    for char,gender in duplicate:
        new_char = find_list_gender(graph,char,gender)
        #print(new_char)
        if new_char[0] in remaining: remaining.append((new_char[1],gender))
        else: remaining.append((new_char[0],gender))
    if not len(char_list_gender) == len(set(char_list_gender)) and count<=10:
        return remove_duplicates(remaining,count = count+1)
    result = list(map(lambda x : x[0], remaining))
    return result
         


def finalize (char_list_prob): # (char,true) gives best prediction
    if all ([x[1] for x in char_list_prob]):
        return char_list_prob
    if len(char_list_prob) == 1: return [char_list_prob[0]]
    if len(char_list_prob) == 2: return list(map(lambda x : x[0] ,char_list_prob))
    test = list()
    predictions = list()
    removed = None
    for char,good in char_list_prob:
        test = list(map(lambda x:x[0],char_list_prob))
        if not good:
            test.remove(char)
            removed = char
            continue
    low = 10.0
    good = dict()
    added = None
    for c in unique:
        char_list = list(map(lambda x: x,test))
        char_list.append(c)
        std = find_std(char_list)
        if std and std < low:
            low = std
            good[low] = char_list
            added = c
    char_list_prob.remove((removed,False))
    char_list_prob.append((added,True))
    #print([x[1] for x in char_list_prob])
    if False in [x[1] for x in char_list_prob]:
        #print ('recurssion')
        finalize(char_list_prob)
    return char_list_prob
                
def get_twins (char_list):
    if 'Lava' in char_list: char_list.append('Kusha')
    elif 'Kusha' in char_list: char_list.append("Lava")
    if 'Vali' in char_list: char_list.append("Sugriva")
    elif 'Sugriva' in char_list: char_list('Vali') 
    return char_list           

