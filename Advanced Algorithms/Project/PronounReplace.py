from docx import *
import time
import numpy as np
from scipy.stats import zscore
import itertools
import collections
plot = 1
start = time.time()
document = Document("D:\\Users\\Hrishi\\Documents\\PESU\\5th Semester\\Advanced Algorithms\\Project\\Ramayana.docx")
bolds=[]
for para in document.paragraphs:
    for run in para.runs:
        if run.bold :
            word=run.text
            if word!="":
                if word[0].isupper():
                    
                    bolds.append(word)

    
sno=1
count=0  
character=dict()
unique=set(bolds)

for i in unique:
    for j in bolds:
        if i==j:
            count+=1
    if i!="":
        character[i]=count 
    count=0;
    
ranked_list=sorted(character, key=lambda x: character[x])[::-1]
for i in ranked_list:
    
    sno+=1

document = Document("D:\\Users\\Hrishi\\Documents\\PESU\\5th Semester\\Advanced Algorithms\\Project\\Ramayana.docx")
bolds=[]
for para in document.paragraphs:
    for run in para.runs:
        if run.bold :
            word=run.text
            if word!="":
                if word[0].isupper():
                    
                    bolds.append(word) #COMMENT: Assuming that every word that starts with a capital letter is a proper noun?

count=0  
character=dict()
unique=set(bolds)
#print(unique)   
for i in unique:
    for j in bolds:
        if i==j:
            count+=1
    if i!="":
        character[i]=count
    count=0;
    
bolds=sorted(character, key=lambda x: character[x])[::-1]

male_pronouns=["he","his","He","His","him","Him"]
female_pronouns=["she","hers","her","She","Hers","Her"]

f=open("D:\\Users\\Hrishi\\Documents\\PESU\\5th Semester\\Advanced Algorithms\\Project\\Ramayana.txt","r")
text=f.read()
sentences=text.split(".")
male=dict()
female=dict()
mcount=0
fcount=0
k=0
for i in bolds:
    
    for j in sentences:
        words=j.split()
        for word in words:
            if word==i:   #COMMENT: why compare word and i? number of pronouns in proximity of bold word?
                for gender in words:
                    if gender in male_pronouns:
                        
                        mcount+=1
                    elif gender in female_pronouns:
                        
                        fcount+=1
                            
        

    male[i]=mcount
    female[i]=fcount
    mcount=0
    fcount=0

final=dict()
males=[]
females=[]
ambiguous=[]
for i in bolds:
    
    if male[i]>female[i]:
        final[i]="M"
        males.append(i)
    elif female[i]>male[i]:
        final[i]="F"
        females.append(i)
    else:
        ambiguous.append(i)

monkeys = list()
monkeys = ['Vali','Hanuman','Sugriva','Kabandha']
malelist = [];
for i in males:
    malelist.append(i);
#print(malelist);    
#print("Females are");
#print(females);

##PREPROCESSING DONE

class PrioQueue:
    def __init__(self):
        self.items =[];
    def isEmpty(self):
        return self.items == []
    def enqueue(self, item):
        self.items.insert(0,item)
    def dequeue(self):
        return self.items.pop();
    def size(self):
        return len(self.items);
    def delete(self):
        for i in range(len(self.items)):
            self.items.pop();
            

q = PrioQueue();
q.enqueue(2);
#print(q.items);
q.enqueue(5);
#print(q.items);
q.enqueue(6);
#print(q.items);
q.enqueue(2);
#print(q.items);


malequeue = PrioQueue();
femalequeue = PrioQueue();
malelist.remove("Kaushalya");
females.append("Kaushalya");
#print("Kaushalya" in females);
f = open("D:\\Users\\Hrishi\\Documents\\PESU\\5th Semester\\Advanced Algorithms\\Project\\output.txt", "w");
connectingwords = ['with', 'With'];
document = Document("D:\\Users\\Hrishi\\Documents\\PESU\\5th Semester\\Advanced Algorithms\\Project\\Ramayana.docx")
for para in document.paragraphs:
    for runs in para.runs:
        for text in runs.text.split('.'):
            #malequeue.delete();
            #femalequeue.delete();
            #print(text);
            for word in text.split(' '):
                
                    
                if(word in malelist):
                    if(word in malequeue.items):
                        malequeue.items.remove(word);
                    malequeue.enqueue(word);
                if(word in females):
                    if(word in femalequeue.items):
                        femalequeue.items.remove(word);
                    femalequeue.enqueue(word);
                if(word in male_pronouns and not malequeue.isEmpty()):
                    
                    f.write(text);
                    f.write("\n");
                    
                    #print("IN SENTENCE");
                    #print(text);
                    f.write(word + " is Replaced With " + malequeue.items[0]);
                    f.write("\n");
                if(word in female_pronouns and not femalequeue.isEmpty()):
                    
                    f.write(text);
                    f.write("\n");
                    f.write(word + " is Replaced With " + femalequeue.items[0]);
                    f.write("\n");
                    

f.close();
